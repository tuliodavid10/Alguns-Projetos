from docx import Document

documento = Document()
documento.add_heading('Relatório de Vendas', 0)
documento.add_paragraph(
    'Neste mês de agosto foram realizadas um total de 10 vendas de notebooks. Segue em anexo a tabela com os dados de cada venda realizada.')
documento.add_heading('Vendas de Agosto', 1)

registros = [
    ['Modelo 1', 'R$1500.00', '25/08/2022', 'Robert', 'Loja 1'],
    ['Modelo 2', 'R$2500.00', '25/08/2022', 'Amanda', 'Loja 1'],
    ['Modelo 2', 'R$2500.00', '25/08/2022', 'Robert', 'Loja 1'],
    ['Modelo 3', 'R$3500.00', '25/08/2022', 'Robert', 'Loja 1'],
    ['Modelo 3', 'R$3500.00', '25/08/2022', 'Carol', 'Loja 1'],
    ['Modelo 3', 'R$3500.00', '25/08/2022', 'Amanda', 'Loja 1'],
    ['Modelo 4', 'R$5000.00', '25/08/2022', 'Robert', 'Loja 1'],
    ['Modelo 4', 'R$5000.00', '25/08/2022', 'Amanda', 'Loja 1'],
    ['Modelo 4', 'R$5000.00', '25/08/2022', 'Robert', 'Loja 1'],
    ['Modelo 4', 'R$5000.00', '25/08/2022', 'Amanda', 'Loja 1'],
]

tabela = documento.add_table(rows=1, cols=5)

cabecalho = tabela.rows[0].cells
cabecalho[0].text = 'Modelo'
cabecalho[1].text = 'Preço'
cabecalho[2].text = 'Data'
cabecalho[3].text = 'Vendedor'
cabecalho[4].text = 'Loja'

for modelo, preco, data, vendedor, loja in registros:
    linha_atual = tabela.add_row().cells
    linha_atual[0].text = modelo
    linha_atual[1].text = preco
    linha_atual[2].text = data
    linha_atual[3].text = vendedor
    linha_atual[4].text = loja

documento.add_paragraph('')
documento.add_paragraph(
    'Para as vendas deste mês, o funcionário Robert foi o funcionário com a maior quantidade de vendas diretas.')
documento.add_paragraph('O lucro total gerado na loja 1 foi de R$37000.00')
documento.save('demo.docx')